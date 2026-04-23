from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x3A, 0x68)

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('BINDING DEBUG TEMPLATE')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

doc.add_paragraph()

h = doc.add_heading('ORIGINAL ATTRIBUTES (should all fill)', level=1)
h.runs[0].font.color.rgb = NAVY

def row(label, token):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(token)

row('Account: ', '[{{Account}}]')
row('Status: ', '[{{Status}}]')
row('Description: ', '[{{Description}}]')
row('SourceOrderId: ', '[{{SourceOrderId}}]')
row('SourceQuoteId: ', '[{{SourceQuoteId}}]')
row('SourceOpportunityId: ', '[{{SourceOpportunityId}}]')
row('IsAssociatedWithClm: ', '[{{IsAssociatedWithClm}}]')

doc.add_paragraph()

h = doc.add_heading('NEW ATTRIBUTES (testing)', level=1)
h.runs[0].font.color.rgb = NAVY
row('ContractNumber: ', '[{{ContractNumber}}]')
row('StartDate: ', '[{{StartDate}}]')
row('EndDate: ', '[{{EndDate}}]')
row('ContractTerm: ', '[{{ContractTerm}}]')
row('CustomerSignedTitle: ', '[{{CustomerSignedTitle}}]')
row('CustomerSignedDate: ', '[{{CustomerSignedDate}}]')
row('CompanySignedDate: ', '[{{CompanySignedDate}}]')
row('AccountName: ', '[{{AccountName}}]')
row('OwnerName: ', '[{{OwnerName}}]')

doc.add_paragraph()

h = doc.add_heading('LINE ITEMS LOOP', level=1)
h.runs[0].font.color.rgb = NAVY
table = doc.add_table(rows=2, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Name'
hdr[1].text = 'Qty'
hdr[2].text = 'UnitPrice'
hdr[3].text = 'TotalPrice'
data = table.rows[1].cells
data[0].text = '{{#SalesContractLineCache}}{{Name}}'
data[1].text = '{{Quantity}}'
data[2].text = '{{UnitPrice}}'
data[3].text = '{{TotalPrice}}{{/SalesContractLineCache}}'

doc.save('StandardContractTemplate.docx')
print('OK - debug template saved')
