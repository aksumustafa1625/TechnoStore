from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

NAVY = RGBColor(0x1F, 0x3A, 0x68)
GREY_TXT = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)


doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def add_bold_line(label, token):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(token)
    return p


# ═══ TITLE ═══
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SALES AGREEMENT')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('Contract No. {{ContractNumber}}')
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = GREY_TXT

doc.add_paragraph()

# ═══ PREAMBLE ═══
pre = doc.add_paragraph()
pre.add_run(
    'This Sales Agreement (the "Agreement") is entered into by and between '
)
r = pre.add_run('TechnoStore Inc. ')
r.bold = True
pre.add_run('(hereinafter "Supplier") and ')
r = pre.add_run('{{AccountName}}')
r.bold = True
pre.add_run(
    ' (hereinafter "Customer"), effective as of {{StartDate}} (the "Effective Date"). '
    'This Agreement sets forth the terms and conditions under which Supplier will '
    'provide products and services to Customer.'
)

doc.add_paragraph()

# ═══ 1. PARTIES ═══
add_heading('1. PARTIES')
add_bold_line('Supplier: ',
              'TechnoStore Inc., a duly organized entity operating in the technology retail and services sector.')
add_bold_line('Customer: ', '{{AccountName}}')
add_bold_line('Contract Owner (Supplier Side): ', '{{OwnerName}}')

# ═══ 2. RECITALS ═══
add_heading('2. RECITALS')
doc.add_paragraph(
    'WHEREAS, Supplier is engaged in the business of selling technology products '
    'and providing related services;'
)
doc.add_paragraph(
    'WHEREAS, Customer desires to purchase such products and services from Supplier '
    'pursuant to the terms set forth in this Agreement;'
)
doc.add_paragraph(
    'NOW, THEREFORE, in consideration of the mutual covenants contained herein, '
    'the parties agree as follows.'
)

# ═══ 3. TERM ═══
add_heading('3. TERM')
add_bold_line('Effective Date: ', '{{StartDate}}')
add_bold_line('Expiration Date: ', '{{EndDate}}')
add_bold_line('Contract Term: ', '{{ContractTerm}} months')
add_bold_line('Contract Status: ', '{{Status}}')
doc.add_paragraph(
    'This Agreement shall commence on the Effective Date and remain in effect until '
    'the Expiration Date, unless terminated earlier in accordance with Section 8.'
)

# ═══ 4. PRODUCTS & SERVICES ═══
add_heading('4. PRODUCTS & SERVICES')
doc.add_paragraph(
    'The following products and services are covered by this Agreement:'
)

# Product table with iteration markers
# First try: Handlebars loop inside table row
# This table has header + one data row marked with {{#SalesContractLineCache}} ... {{/SalesContractLineCache}}
table = doc.add_table(rows=2, cols=4)
table.style = 'Light Grid Accent 1'
table.autofit = False

widths = [Inches(3.0), Inches(0.8), Inches(1.2), Inches(1.5)]
for row in table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = widths[idx]

hdr = table.rows[0].cells
hdr[0].text = 'Product / Service'
hdr[1].text = 'Qty'
hdr[2].text = 'Unit Price'
hdr[3].text = 'Total'

for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

# Data row with iteration tokens — use {{Name}} instead of {{Product}} (Name = product name)
data = table.rows[1].cells
data[0].text = '{{#SalesContractLineCache}}{{Name}}'
data[1].text = '{{Quantity}}'
data[2].text = '{{UnitPrice}}'
data[3].text = '{{TotalPrice}}{{/SalesContractLineCache}}'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Total amounts shown in the contract currency.')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = LIGHT_GREY

# ═══ 5. PAYMENT TERMS ═══
add_heading('5. PAYMENT TERMS')
doc.add_paragraph(
    'Customer shall pay all amounts due within thirty (30) days of the invoice date. '
    'Invoices will be issued pursuant to the payment schedule in the source order. '
    'Late payments shall accrue interest at the rate of 1.5% per month or the maximum '
    'rate permitted by applicable law, whichever is lower.'
)

# ═══ 6. CONFIDENTIALITY ═══
add_heading('6. CONFIDENTIALITY')
doc.add_paragraph(
    'Each party agrees to keep confidential all non-public information disclosed by the '
    'other party in connection with this Agreement. Confidential Information shall not be '
    'disclosed to any third party without prior written consent, and shall be used solely '
    'for the purposes of performing this Agreement.'
)

# ═══ 7. WARRANTY & LIMITATION ═══
add_heading('7. WARRANTY & LIMITATION OF LIABILITY')
doc.add_paragraph(
    'Supplier warrants that the products and services will conform to the specifications '
    "in the source order. Supplier's aggregate liability under this Agreement shall not "
    'exceed the total fees paid by Customer in the twelve (12) months preceding the claim. '
    'Neither party shall be liable for indirect, incidental, or consequential damages.'
)

# ═══ 8. TERMINATION ═══
add_heading('8. TERMINATION')
doc.add_paragraph(
    'Either party may terminate this Agreement for material breach upon thirty (30) days '
    'written notice if the breach is not cured within that period.'
)

# ═══ 9. GOVERNING LAW ═══
add_heading('9. GOVERNING LAW')
doc.add_paragraph(
    'This Agreement shall be governed by and construed in accordance with the applicable '
    'laws, without regard to conflict-of-law principles. Any disputes arising under this '
    "Agreement shall be resolved in the courts of competent jurisdiction at Supplier's "
    'principal place of business.'
)

# ═══ 10. ENTIRE AGREEMENT ═══
add_heading('10. ENTIRE AGREEMENT')
doc.add_paragraph(
    'This Agreement, together with the referenced source documents, constitutes the entire '
    'agreement between the parties with respect to the subject matter hereof, superseding '
    'all prior communications. Any amendments must be in writing and signed by both parties.'
)

# ═══ 11. SIGNATURES ═══
doc.add_paragraph()
add_heading('11. SIGNATURES')
doc.add_paragraph(
    'IN WITNESS WHEREOF, the parties have executed this Agreement as of the dates set '
    'forth below.'
)

sig = doc.add_table(rows=2, cols=2)
sig.style = 'Light Grid Accent 1'
hdr = sig.rows[0].cells
hdr[0].text = 'SUPPLIER — TechnoStore Inc.'
hdr[1].text = 'CUSTOMER — {{AccountName}}'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

row = sig.rows[1].cells
row[0].text = 'Signed by: {{CompanySignedByName}}\nDate: {{CompanySignedDate}}\n\n\nSignature: ______________________'
row[1].text = 'Signed by: {{CustomerSignedByName}}\nTitle: {{CustomerSignedTitle}}\nDate: {{CustomerSignedDate}}\n\nSignature: ______________________'

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
f = p.add_run('— End of Agreement —')
f.italic = True
f.font.size = Pt(9)
f.font.color.rgb = LIGHT_GREY

doc.save('StandardContractTemplate.docx')
print('OK - contract template with products table saved')
